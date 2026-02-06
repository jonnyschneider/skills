#!/usr/bin/env python3
"""
Create a complete reference template for pandoc markdown-to-docx conversion.

This template includes all standard styles that pandoc expects/uses:
- Headings 1-6
- Normal paragraphs
- Lists (bullet and numbered)
- Code blocks (Source Code) with background shading
- Inline code (Verbatim Char)
- Block quotes (Block Text) with left border
- Tables with header styling
- Hyperlinks

Usage:
    create-reference-template.py <output.docx> [--font-body "Font Name"] [--font-heading "Font Name"] [--font-mono "Font Name"]
"""

import sys
import zipfile
import tempfile
from pathlib import Path
from xml.etree import ElementTree as ET

import click
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def create_element(tag: str, **attribs) -> OxmlElement:
    """Create an OxmlElement with attributes."""
    elem = OxmlElement(tag)
    for key, val in attribs.items():
        elem.set(qn(f'w:{key}'), str(val))
    return elem


def add_shading_to_style(style, fill_color: str):
    """Add background shading to a paragraph style."""
    # Access the style's XML element
    style_elem = style.element

    # Get or create pPr (paragraph properties)
    pPr = style_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style_elem.insert(0, pPr)

    # Remove existing shading if any
    existing_shd = pPr.find(qn('w:shd'))
    if existing_shd is not None:
        pPr.remove(existing_shd)

    # Add shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)


def add_borders_to_style(style, color: str = 'DDDDDD', size: int = 4, sides: list = None):
    """Add borders to a paragraph style."""
    if sides is None:
        sides = ['top', 'left', 'bottom', 'right']

    style_elem = style.element

    pPr = style_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style_elem.insert(0, pPr)

    # Remove existing borders
    existing_bdr = pPr.find(qn('w:pBdr'))
    if existing_bdr is not None:
        pPr.remove(existing_bdr)

    # Add borders
    pBdr = OxmlElement('w:pBdr')
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), color)
        pBdr.append(border)

    pPr.append(pBdr)


def lighten_color(hex_color: str, factor: float = 0.7) -> str:
    """Lighten a hex color by mixing with white."""
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)

    # Mix with white
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)

    return f'{r:02X}{g:02X}{b:02X}'


def create_table_style_element(
    style_id: str,
    name: str,
    accent_color: str = '4472C4',
    font_name: str = None,
    font_size_pt: int = None,
    compact: bool = False,
) -> OxmlElement:
    """Create a table style element using OxmlElement (proper namespace handling)."""
    # Calculate colors from accent - dark header with white text
    header_bg = accent_color  # Full accent color for header
    header_text = 'FFFFFF'  # White text on dark header
    alt_row_bg = lighten_color(accent_color, 0.9)  # Very light tint for alternating rows
    border_color = lighten_color(accent_color, 0.4)  # Medium tint for borders

    # Cell margins: compact vs standard
    if compact:
        cell_margins = [('top', '36'), ('left', '72'), ('bottom', '36'), ('right', '72')]  # ~0.025" / 0.05"
    else:
        cell_margins = [('top', '72'), ('left', '115'), ('bottom', '72'), ('right', '115')]  # ~0.05" / 0.08"

    # Create style element
    style = OxmlElement('w:style')
    style.set(qn('w:type'), 'table')
    style.set(qn('w:styleId'), style_id)

    # Name
    name_elem = OxmlElement('w:name')
    name_elem.set(qn('w:val'), name)
    style.append(name_elem)

    # BasedOn
    based_on = OxmlElement('w:basedOn')
    based_on.set(qn('w:val'), 'TableNormal')
    style.append(based_on)

    # Table properties
    tbl_pr = OxmlElement('w:tblPr')

    # Borders
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    # Cell margins
    tbl_cell_mar = OxmlElement('w:tblCellMar')
    for side, val in cell_margins:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), val)
        margin.set(qn('w:type'), 'dxa')
        tbl_cell_mar.append(margin)
    tbl_pr.append(tbl_cell_mar)

    style.append(tbl_pr)

    # Default paragraph properties for all cells (compact spacing for lists)
    pPr_default = OxmlElement('w:pPr')
    spacing_default = OxmlElement('w:spacing')
    spacing_default.set(qn('w:before'), '40')  # ~2pt before
    spacing_default.set(qn('w:after'), '40')   # ~2pt after
    spacing_default.set(qn('w:line'), '240')   # Single line spacing
    spacing_default.set(qn('w:lineRule'), 'auto')
    pPr_default.append(spacing_default)
    # Compact list indentation for table content
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '144')  # 0.1" for lists in tables
    ind.set(qn('w:hanging'), '144')  # Hanging indent for bullets
    pPr_default.append(ind)
    style.append(pPr_default)

    # Default run properties for table text
    rPr_default = OxmlElement('w:rPr')
    if font_name:
        rFonts_default = OxmlElement('w:rFonts')
        rFonts_default.set(qn('w:ascii'), font_name)
        rFonts_default.set(qn('w:hAnsi'), font_name)
        rPr_default.append(rFonts_default)
    style.append(rPr_default)

    # First row styling (header)
    first_row = OxmlElement('w:tblStylePr')
    first_row.set(qn('w:type'), 'firstRow')

    # Paragraph properties for header row
    pPr = OxmlElement('w:pPr')
    # Reduce spacing in header cells
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    first_row.append(pPr)

    # Run properties (text formatting)
    rPr = OxmlElement('w:rPr')
    # Bold
    bold = OxmlElement('w:b')
    rPr.append(bold)
    bold_cs = OxmlElement('w:bCs')  # Bold for complex scripts too
    rPr.append(bold_cs)
    # White text color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), header_text)
    rPr.append(color)
    # Font name (if specified)
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)  # Complex script font
        rPr.append(rFonts)
    # Font size (if specified)
    if font_size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_pt * 2))  # Half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(szCs)
    first_row.append(rPr)

    # Cell properties (background)
    tcPr = OxmlElement('w:tcPr')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), header_bg)
    tcPr.append(shd)
    # Vertical alignment (center)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    first_row.append(tcPr)

    style.append(first_row)

    # Alternating row styling
    band1 = OxmlElement('w:tblStylePr')
    band1.set(qn('w:type'), 'band1Horz')

    tcPr2 = OxmlElement('w:tcPr')
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'), alt_row_bg)
    tcPr2.append(shd2)
    band1.append(tcPr2)

    style.append(band1)

    return style


def inject_compact_numbering(doc_path: Path):
    """Inject compact numbering definitions for bullet and numbered lists."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        # Extract docx
        with zipfile.ZipFile(doc_path) as zf:
            zf.extractall(tmp)

        # Create numbering.xml with compact indents
        # Indent values in twips: 144 = 0.1", 216 = 0.15", 288 = 0.2"
        numbering_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <!-- Compact bullet list -->
    <w:abstractNum w:abstractNumId="0">
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="&#x2022;"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="216" w:hanging="216"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="&#x25E6;"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="432" w:hanging="216"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="&#x25AA;"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="648" w:hanging="216"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    <!-- Compact numbered list -->
    <w:abstractNum w:abstractNumId="1">
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="288" w:hanging="288"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="lowerLetter"/>
            <w:lvlText w:val="%2."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="576" w:hanging="288"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    <w:num w:numId="1">
        <w:abstractNumId w:val="0"/>
    </w:num>
    <w:num w:numId="2">
        <w:abstractNumId w:val="1"/>
    </w:num>
</w:numbering>'''

        numbering_path = tmp / 'word' / 'numbering.xml'
        numbering_path.write_text(numbering_xml)

        # Update Content_Types.xml to include numbering
        ct_path = tmp / '[Content_Types].xml'
        tree = ET.parse(ct_path)
        root = tree.getroot()
        ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
        ET.register_namespace('', ns)

        # Check if numbering override exists
        has_numbering = False
        for override in root.findall(f'{{{ns}}}Override'):
            if 'numbering' in override.get('PartName', ''):
                has_numbering = True
                break

        if not has_numbering:
            override = ET.SubElement(root, f'{{{ns}}}Override')
            override.set('PartName', '/word/numbering.xml')
            override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml')
            tree.write(ct_path, xml_declaration=True, encoding='UTF-8')

        # Update document.xml.rels to reference numbering
        rels_path = tmp / 'word' / '_rels' / 'document.xml.rels'
        tree = ET.parse(rels_path)
        root = tree.getroot()
        ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
        ET.register_namespace('', ns)

        # Check if numbering relationship exists
        has_numbering_rel = False
        max_id = 0
        for rel in root:
            rid = rel.get('Id', '')
            if rid.startswith('rId'):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
            if 'numbering' in rel.get('Target', ''):
                has_numbering_rel = True

        if not has_numbering_rel:
            rel = ET.SubElement(root, f'{{{ns}}}Relationship')
            rel.set('Id', f'rId{max_id + 1}')
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
            rel.set('Target', 'numbering.xml')
            tree.write(rels_path, xml_declaration=True, encoding='UTF-8')

        # Repack
        doc_path.unlink()
        with zipfile.ZipFile(doc_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for f in tmp.rglob('*'):
                if f.is_file():
                    zf.write(f, f.relative_to(tmp))


def inject_table_style(doc_path: Path, accent_color: str = '4472C4', font_name: str = None):
    """Inject table styles into the document using python-docx internals."""
    from docx import Document

    # Open the document
    doc = Document(doc_path)

    # Access the styles part
    styles_element = doc.styles.element

    # Remove existing Table and Compact Table styles if they exist
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for style_id in ['Table', 'CompactTable']:
        for style in styles_element.findall('.//w:style', ns):
            if style.get(qn('w:styleId')) == style_id:
                styles_element.remove(style)
                break

    # Create and add the standard Table style
    table_style = create_table_style_element(
        'Table', 'Table',
        accent_color=accent_color,
        font_name=font_name,
    )
    styles_element.append(table_style)

    # Create and add the Compact Table style (tighter margins, smaller text)
    compact_table_style = create_table_style_element(
        'CompactTable', 'Compact Table',
        accent_color=accent_color,
        font_name=font_name,
        font_size_pt=9,
        compact=True,
    )
    styles_element.append(compact_table_style)

    # Save
    doc.save(doc_path)


def fix_doc_defaults_fonts(doc_path: Path, font_body: str, font_heading: str = None):
    """Fix docDefaults and theme fonts to use explicit fonts.

    python-docx sets fonts via theme references (minorHAnsi, majorHAnsi) which
    resolve to Calibri. This function:
    1. Updates docDefaults to use explicit font names
    2. Updates theme1.xml to set the theme fonts to our desired fonts
    """
    from lxml import etree

    if font_heading is None:
        font_heading = font_body

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        with zipfile.ZipFile(doc_path) as zf:
            zf.extractall(tmp)

        # Fix styles.xml docDefaults
        styles_xml = tmp / 'word' / 'styles.xml'
        if styles_xml.exists():
            parser = etree.XMLParser(remove_blank_text=False)
            tree = etree.parse(str(styles_xml), parser)
            root = tree.getroot()

            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            w_ns = f'{{{ns["w"]}}}'

            # Find docDefaults > rPrDefault > rPr > rFonts
            doc_defaults = root.find('.//w:docDefaults', ns)
            if doc_defaults is not None:
                rFonts = doc_defaults.find('.//w:rPrDefault/w:rPr/w:rFonts', ns)
                if rFonts is not None:
                    # Remove theme attributes
                    for attr in ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']:
                        full_attr = f'{w_ns}{attr}'
                        if full_attr in rFonts.attrib:
                            del rFonts.attrib[full_attr]

                    # Set explicit font names
                    rFonts.set(f'{w_ns}ascii', font_body)
                    rFonts.set(f'{w_ns}hAnsi', font_body)
                    rFonts.set(f'{w_ns}eastAsia', font_body)
                    rFonts.set(f'{w_ns}cs', font_body)

            tree.write(str(styles_xml), xml_declaration=True, encoding='UTF-8', standalone=True)

        # Fix theme1.xml font scheme
        theme_xml = tmp / 'word' / 'theme' / 'theme1.xml'
        if theme_xml.exists():
            parser = etree.XMLParser(remove_blank_text=False)
            tree = etree.parse(str(theme_xml), parser)
            root = tree.getroot()

            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'

            # Update majorFont (headings)
            for latin in root.findall(f'.//{{{a_ns}}}majorFont/{{{a_ns}}}latin'):
                latin.set('typeface', font_heading)

            # Update minorFont (body)
            for latin in root.findall(f'.//{{{a_ns}}}minorFont/{{{a_ns}}}latin'):
                latin.set('typeface', font_body)

            tree.write(str(theme_xml), xml_declaration=True, encoding='UTF-8', standalone=True)

        # Repack
        doc_path.unlink()
        with zipfile.ZipFile(doc_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for f in tmp.rglob('*'):
                if f.is_file():
                    zf.write(f, f.relative_to(tmp))


def create_reference_template(
    output_path: str,
    font_body: str = 'Calibri',
    font_heading: str = 'Calibri',
    font_mono: str = 'Consolas',
    accent_color: str = '4472C4',  # Blue
    heading_color: str = None,  # Hex color for headings (defaults to black)
    include_table_style: bool = True,  # Toggle for debugging
) -> dict:
    """
    Create a complete reference template.

    Args:
        output_path: Path for output DOCX
        font_body: Font for body text
        font_heading: Font for headings
        font_mono: Font for code
        accent_color: Hex color for accents (tables, etc)
        heading_color: Hex color for headings (default: black)

    Returns:
        Dictionary with template info
    """
    # Parse heading color
    if heading_color:
        heading_rgb = RGBColor(
            int(heading_color[0:2], 16),
            int(heading_color[2:4], 16),
            int(heading_color[4:6], 16)
        )
    else:
        heading_rgb = RGBColor(0x00, 0x00, 0x00)  # Black
    doc = Document()
    styles_created = []

    # =========================================
    # PARAGRAPH STYLES
    # =========================================

    # Normal (body text)
    normal = doc.styles['Normal']
    normal.font.name = font_body
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    styles_created.append('Normal')

    # Headings 1-6
    heading_config = [
        (1, 24, 24, 8),   # level, size, space_before, space_after
        (2, 18, 18, 6),
        (3, 14, 14, 4),
        (4, 12, 12, 4),
        (5, 11, 10, 2),
        (6, 11, 10, 2),
    ]

    for level, size, space_before, space_after in heading_config:
        style = doc.styles[f'Heading {level}']
        style.font.name = font_heading
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = heading_rgb
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True
        styles_created.append(f'Heading {level}')

    # Title
    title = doc.styles['Title']
    title.font.name = font_heading
    title.font.size = Pt(28)
    title.font.bold = True
    title.paragraph_format.space_after = Pt(12)
    styles_created.append('Title')

    # Subtitle
    subtitle = doc.styles['Subtitle']
    subtitle.font.name = font_heading
    subtitle.font.size = Pt(16)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    styles_created.append('Subtitle')

    # Block Text (for blockquotes) - with left border
    try:
        block_text = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        block_text = doc.styles['Block Text']
    block_text.font.name = font_body
    block_text.font.size = Pt(11)
    block_text.font.italic = True
    block_text.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    block_text.paragraph_format.left_indent = Inches(0.4)
    block_text.paragraph_format.right_indent = Inches(0.2)
    block_text.paragraph_format.space_before = Pt(8)
    block_text.paragraph_format.space_after = Pt(8)
    # Add left border (quote style)
    add_borders_to_style(block_text, color=accent_color, size=18, sides=['left'])
    styles_created.append('Block Text')

    # Source Code (for code blocks) - with background and border
    try:
        source_code = doc.styles.add_style('Source Code', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        source_code = doc.styles['Source Code']
    source_code.font.name = font_mono
    source_code.font.size = Pt(9)
    source_code.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    source_code.paragraph_format.space_before = Pt(6)
    source_code.paragraph_format.space_after = Pt(6)
    source_code.paragraph_format.left_indent = Inches(0.15)
    source_code.paragraph_format.right_indent = Inches(0.15)
    # Add background shading
    add_shading_to_style(source_code, 'F5F5F5')
    # Add subtle border
    add_borders_to_style(source_code, color='E0E0E0', size=4)
    styles_created.append('Source Code')

    # List Paragraph (for bullet/numbered lists)
    # Minimal indentation for tables - bullet hangs, text aligns to edge
    list_para = doc.styles['List Paragraph']
    list_para.font.name = font_body
    list_para.font.size = Pt(11)
    list_para.paragraph_format.left_indent = Inches(0.15)  # Minimal for tables
    list_para.paragraph_format.first_line_indent = Inches(-0.15)  # Hanging indent for bullet
    list_para.paragraph_format.space_after = Pt(2)  # Tighter spacing
    styles_created.append('List Paragraph')

    # Compact List (alternative with minimal indentation for tables)
    try:
        compact_list = doc.styles.add_style('Compact List', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        compact_list = doc.styles['Compact List']
    compact_list.base_style = doc.styles['List Paragraph']
    compact_list.font.name = font_body
    compact_list.font.size = Pt(10)
    compact_list.paragraph_format.left_indent = Inches(0.1)
    compact_list.paragraph_format.space_after = Pt(1)
    compact_list.paragraph_format.space_before = Pt(0)
    styles_created.append('Compact List')

    # =========================================
    # CHARACTER STYLES
    # =========================================

    # Verbatim Char (for inline code) - with subtle background
    try:
        verbatim = doc.styles.add_style('Verbatim Char', WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        verbatim = doc.styles['Verbatim Char']
    verbatim.font.name = font_mono
    verbatim.font.size = Pt(10)
    verbatim.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)  # Reddish for inline code
    # Note: Character styles can't have background in python-docx easily
    styles_created.append('Verbatim Char')

    # Strong (bold)
    strong = doc.styles['Strong']
    strong.font.bold = True
    styles_created.append('Strong')

    # Emphasis (italic)
    emphasis = doc.styles['Emphasis']
    emphasis.font.italic = True
    styles_created.append('Emphasis')

    # Hyperlink
    try:
        hyperlink = doc.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        hyperlink = doc.styles['Hyperlink']
    hyperlink.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    hyperlink.font.underline = True
    styles_created.append('Hyperlink')

    # Figure/Image Caption (pandoc uses this for image alt text)
    try:
        caption = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        caption = doc.styles['Caption']
    caption.font.name = font_body
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles_created.append('Caption')

    # Image Paragraph (for standalone images)
    try:
        image_para = doc.styles.add_style('Image', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        image_para = doc.styles['Image']
    image_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_before = Pt(12)
    image_para.paragraph_format.space_after = Pt(6)
    styles_created.append('Image')

    # =========================================
    # SAVE AND POST-PROCESS
    # =========================================

    # Add placeholder paragraph
    doc.add_paragraph()

    # Save initial document
    output_path = Path(output_path)
    doc.save(output_path)

    # Inject compact numbering definitions for lists
    inject_compact_numbering(output_path)
    styles_created.append('Compact numbering (bullets/numbers)')

    # Inject table styles via XML (python-docx table style support is limited)
    if include_table_style:
        inject_table_style(output_path, accent_color=accent_color, font_name=font_body)
        styles_created.append('Table')
        styles_created.append('Compact Table')

    # Fix docDefaults and theme fonts to use explicit fonts instead of theme references
    fix_doc_defaults_fonts(output_path, font_body, font_heading)

    return {
        'output': str(output_path),
        'fonts': {
            'body': font_body,
            'heading': font_heading,
            'mono': font_mono,
        },
        'accent_color': accent_color,
        'styles': styles_created,
    }


@click.command()
@click.argument('output', type=click.Path())
@click.option('--font-body', default='Calibri', help='Body text font (default: Calibri)')
@click.option('--font-heading', default='Calibri', help='Heading font (default: Calibri)')
@click.option('--font-mono', default='Consolas', help='Monospace font (default: Consolas)')
@click.option('--accent-color', default='4472C4', help='Accent color hex (default: 4472C4 blue)')
@click.option('--heading-color', default=None, help='Heading color hex (default: black)')
@click.option('--no-table-style', is_flag=True, help='Skip table style injection (for debugging)')
def main(output, font_body, font_heading, font_mono, accent_color, heading_color, no_table_style):
    """Create a reference template for pandoc markdown conversion."""

    print(f"Creating reference template: {output}")
    print(f"  Body font: {font_body}")
    print(f"  Heading font: {font_heading}")
    print(f"  Mono font: {font_mono}")
    print(f"  Accent color: #{accent_color}")
    print(f"  Heading color: #{heading_color if heading_color else '000000 (black)'}")
    print(f"  Table style: {'skipped' if no_table_style else 'included'}")

    result = create_reference_template(
        output,
        font_body=font_body,
        font_heading=font_heading,
        font_mono=font_mono,
        accent_color=accent_color,
        heading_color=heading_color,
        include_table_style=not no_table_style,
    )

    print(f"\nStyles configured ({len(result['styles'])}):")
    for style in result['styles']:
        print(f"  - {style}")

    print(f"\nTemplate created: {result['output']}")
    print("\nUsage with pandoc:")
    print(f"  pandoc input.md --reference-doc={output} -o output.docx")


if __name__ == '__main__':
    main()
